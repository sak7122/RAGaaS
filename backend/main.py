from __future__ import annotations

import hashlib
import hmac
import io
import json
import logging
import os
import re
import time
import urllib.request
from datetime import datetime, timezone
from pathlib import Path
from typing import Annotated

from fastapi import BackgroundTasks, Depends, FastAPI, File, Header, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document
from starlette.middleware.base import BaseHTTPMiddleware

from backend.config import config
from backend.firebase_services import (
    MemberRecord,
    Principal,
    QUERY_LIMIT,
    create_member_store,
    set_tenant_claim,
    create_usage_store,
    verify_firebase_token,
)
from backend.storage import create_storage_backend
from backend.index_store import create_index_store
from backend.rag import create_embedder, create_generator
from backend.insights import create_insights_store
from backend.tenant_profile import create_tenant_profile_store, prettify
from backend.mailer import send_invite_email, smtp_configured
from backend.share_store import create_share_store
from backend.widget_store import create_widget_key_store
from backend.slack_store import create_slack_store

# ── Structured logging ────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format='{"time":"%(asctime)s","level":"%(levelname)s","msg":"%(message)s"}',
    datefmt="%Y-%m-%dT%H:%M:%SZ",
)
log = logging.getLogger("ragaas")

# ── Constants ─────────────────────────────────────────────────────────────────
DATA_DIR       = Path("local_data")
UPLOAD_DIR     = DATA_DIR / "uploads"
INDEX_FILE     = DATA_DIR / "index.json"
MAX_UPLOAD_BYTES = 50 * 1024 * 1024
PDF_MAGIC      = b"%PDF"
ZIP_MAGIC      = b"PK\x03\x04"  # docx is a zip container — disambiguate by extension
VALID_ROLES    = {"admin", "uploader", "viewer"}
# Retrieval engine label surfaced in the UI. Accurate for the DIY hybrid path;
# set RAG_ENGINE_LABEL="Vertex AI Search" if you migrate to managed retrieval.
RAG_ENGINE_LABEL = os.getenv("RAG_ENGINE_LABEL", "Hybrid Vector Search")
HYBRID_VECTOR_WEIGHT = 0.7   # blend: 0.7*cosine + 0.3*keyword overlap
SLACK_SIGNING_SECRET = os.getenv("SLACK_SIGNING_SECRET", "")

# Dev-only mock tokens — empty dict in production
TOKEN_TENANT_MAP: dict[str, str] = (
    {
        "mock-tenant-token-abc": "tenant-demo",
        "tenant-a-token":        "tenant-a",
        "tenant-b-token":        "tenant-b",
    }
    if config.env == "development" else {}
)

if config.env == "development":
    log.info("dev mode: mock token map active (%d tokens)", len(TOKEN_TENANT_MAP))

# ── Widget / Share CORS ───────────────────────────────────────────────────────
# Widget routes and public share reads must be accessible from any origin
# (widget embeds on 3rd-party sites). Widget keys replace browser credentials so
# allow_credentials is NOT set, which is compatible with Access-Control-Allow-Origin: *.
class WidgetCORSMiddleware(BaseHTTPMiddleware):
    _WIDGET_PATHS = ("/api/widget/", "/api/share/")

    async def dispatch(self, request: Request, call_next):
        is_widget = any(request.url.path.startswith(p) for p in self._WIDGET_PATHS)
        if is_widget and request.method == "OPTIONS":
            return JSONResponse(
                {},
                headers={
                    "Access-Control-Allow-Origin": "*",
                    "Access-Control-Allow-Methods": "GET, POST, DELETE, OPTIONS",
                    "Access-Control-Allow-Headers": "Content-Type, X-Widget-Key",
                },
            )
        response = await call_next(request)
        if is_widget:
            response.headers["Access-Control-Allow-Origin"] = "*"
        return response


# ── App ───────────────────────────────────────────────────────────────────────
app = FastAPI(title="RAGaaS API", version="0.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=config.cors_origin_regex,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
)
app.add_middleware(WidgetCORSMiddleware)  # outer layer — handles widget/share/slack paths

usage_store          = create_usage_store()
member_store         = create_member_store()
storage              = create_storage_backend(config.env, UPLOAD_DIR, config.gcs_bucket)
index_store          = create_index_store(config.env, INDEX_FILE)
embedder             = create_embedder()
generator            = create_generator()
insights_store       = create_insights_store(config.env)
tenant_profile_store = create_tenant_profile_store(config.env)
share_store          = create_share_store(config.env)
widget_key_store     = create_widget_key_store(config.env)
slack_store          = create_slack_store(config.env)
log.info("startup env=%s emulator=%s embedder=%s generator=%s",
         config.env, config.use_emulator, type(embedder).__name__, type(generator).__name__)


# ── Models ────────────────────────────────────────────────────────────────────
class TurnHistory(BaseModel):
    role: str   # "user" or "assistant"
    text: str = Field(max_length=2000)


class ChatRequest(BaseModel):
    message: str = Field(min_length=1, max_length=4000)
    history: list[TurnHistory] = Field(default_factory=list, max_length=4)


class Citation(BaseModel):
    file_name: str
    page: int
    chunk_index: int
    excerpt: str
    score: float = 0.0          # normalized relevance 0..1


class RetrievalTrace(BaseModel):
    engine: str                 # retrieval engine label
    query_terms: list[str]      # terms extracted from the question
    chunks_searched: int        # total chunks scanned for this tenant
    candidates_ranked: int      # chunks with a non-zero score
    top_k: int                  # how many returned
    max_score: float            # best relevance score
    latency_ms: int             # retrieval time


class ChatResponse(BaseModel):
    answer: str
    citations: list[Citation]
    retrieval: RetrievalTrace
    tenant_id: str
    queries_used: int
    query_limit: int


class TenantStatus(BaseModel):
    tenant_id: str
    tenant_name: str
    queries_used: int
    query_limit: int
    documents: int


class TenantProfileRequest(BaseModel):
    name: str = Field(min_length=1, max_length=80)


class DocumentMeta(BaseModel):
    file_name: str
    chunks: int
    uploaded_at: str


class QuestionStat(BaseModel):
    question: str
    count: int
    avg_score: float


class KnowledgeGap(BaseModel):
    question: str
    count: int
    best_score: float
    avg_score: float


class InsightsResponse(BaseModel):
    total_queries: int
    avg_confidence: float
    answered_rate: float
    top_questions: list[QuestionStat]
    gaps: list[KnowledgeGap]
    window: int


class MemberOut(BaseModel):
    uid: str
    email: str
    role: str
    invited_at: str
    joined_at: str | None


class InviteRequest(BaseModel):
    email: str = Field(min_length=3, max_length=254)
    role: str = Field(default="viewer")


class AcceptInviteRequest(BaseModel):
    tenant_id: str = Field(min_length=1, max_length=128)
    invite_uid: str | None = None  # informational; authorization is by email


class RoleRequest(BaseModel):
    role: str


class ShareCreateRequest(BaseModel):
    question: str = Field(min_length=1, max_length=4000)
    answer: str = Field(min_length=1, max_length=20000)
    citations: list[dict] = Field(default_factory=list)


class WidgetKeyCreateRequest(BaseModel):
    label: str = Field(min_length=1, max_length=80, default="My website")


class SlackConnectRequest(BaseModel):
    team_id: str = Field(min_length=1, max_length=32)
    team_name: str = Field(min_length=1, max_length=128)


# ── Auth & role guards ────────────────────────────────────────────────────────
def principal_from_auth(
    authorization: Annotated[str | None, Header()] = None,
    x_widget_key: Annotated[str | None, Header()] = None,
) -> Principal:
    # Widget API key auth — viewer-only, no Firebase token needed
    if x_widget_key:
        tenant_id = widget_key_store.resolve_tenant(x_widget_key)
        if not tenant_id:
            raise HTTPException(status_code=401, detail="Invalid widget key")
        return Principal(uid="widget-user", tenant_id=tenant_id, email=None, role="viewer")
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")
    token = authorization.removeprefix("Bearer ").strip()
    tenant_id = TOKEN_TENANT_MAP.get(token)
    if tenant_id:
        # Dev mock tokens always get admin so all endpoints are testable
        return Principal(uid=f"dev-{tenant_id}", tenant_id=tenant_id,
                         email=f"{tenant_id}@ragaas.local", role="admin")
    return verify_firebase_token(token, member_store)


def require_role(principal: Principal, *roles: str) -> None:
    if principal.role not in roles:
        raise HTTPException(
            status_code=403,
            detail=f"Role '{principal.role}' cannot perform this action. Required: {list(roles)}",
        )


# load_index / save_index kept for test compatibility only (dev local store)
def load_index() -> dict:
    if not INDEX_FILE.exists():
        return {"documents": []}
    return json.loads(INDEX_FILE.read_text(encoding="utf-8"))


def save_index(index: dict) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    INDEX_FILE.write_text(json.dumps(index, indent=2), encoding="utf-8")


# ── Document helpers ──────────────────────────────────────────────────────────
def safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", name).strip("._")
    return cleaned or "upload.pdf"


def assert_safe_path(resolved: Path, base: Path) -> None:
    try:
        resolved.relative_to(base.resolve())
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid filename")


def chunk_text(text: str, chunk_size: int = 300, overlap: int = 50) -> list[str]:
    words = text.split()
    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start += chunk_size - overlap
    return chunks or [""]


def _extract_pdf_chunks(content: bytes) -> list[dict]:
    try:
        reader = PdfReader(io.BytesIO(content))
        result = []
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            for idx, chunk in enumerate(chunk_text(page_text)):
                if chunk.strip():               # skip empty chunks (blank/image pages)
                    result.append({"page": page_num, "chunk_index": idx, "text": chunk})
        # No selectable text -> almost certainly a scanned/image PDF (needs OCR).
        if not result:
            raise HTTPException(
                status_code=422,
                detail=("No selectable text found in this PDF. It looks like a scanned or "
                        "image-only document — OCR isn't supported yet. Upload a text-based PDF."),
            )
        return result
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("PDF parse failed: %s", exc)
        raise HTTPException(
            status_code=400,
            detail="Could not parse PDF — file may be corrupted or encrypted",
        )


def _extract_docx_chunks(content: bytes) -> list[dict]:
    try:
        doc = Document(io.BytesIO(content))
        # Paragraphs + table cells — docx has no reliable page boundaries, so page=1.
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        full_text = "\n".join(t for t in parts if t and t.strip())
        result = [
            {"page": 1, "chunk_index": idx, "text": chunk}
            for idx, chunk in enumerate(chunk_text(full_text))
            if chunk.strip()
        ]
        if not result:
            raise HTTPException(
                status_code=422,
                detail="No text found in this Word document. Upload a document that contains text.",
            )
        return result
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("DOCX parse failed: %s", exc)
        raise HTTPException(
            status_code=400,
            detail="Could not parse Word document — file may be corrupted. Legacy .doc isn't supported; save as .docx.",
        )


def extract_chunks(content: bytes, file_name: str) -> list[dict]:
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    if content[:4] == PDF_MAGIC:
        return _extract_pdf_chunks(content)
    if ext == "docx" and content[:4] == ZIP_MAGIC:
        return _extract_docx_chunks(content)
    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Upload a PDF or a Word (.docx) document.",
    )


def tokenize(text: str) -> set[str]:
    return {w for w in re.findall(r"[a-z0-9]{3,}", text.lower())}


def enforce_quota(tenant_id: str) -> int:
    return usage_store.increment_or_reject(tenant_id)


# ── Routes: health ────────────────────────────────────────────────────────────
@app.get("/api/health")
def health() -> dict:
    return {"ok": True, "env": config.env, "time": datetime.now(timezone.utc).isoformat()}


# ── Routes: tenant status ─────────────────────────────────────────────────────
@app.get("/api/tenant/status", response_model=TenantStatus)
def tenant_status(principal: Annotated[Principal, Depends(principal_from_auth)]) -> TenantStatus:
    tenant_id = principal.tenant_id
    docs = index_store.list_docs(tenant_id)
    name = tenant_profile_store.get_name(tenant_id) or prettify(tenant_id)
    return TenantStatus(
        tenant_id=tenant_id,
        tenant_name=name,
        queries_used=usage_store.get_count(tenant_id),
        query_limit=QUERY_LIMIT,
        documents=len(docs),
    )


@app.put("/api/tenant/profile")
def set_tenant_profile(
    body: TenantProfileRequest,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    require_role(principal, "admin")
    tenant_profile_store.set_name(principal.tenant_id, body.name.strip())
    log.info("tenant_name set tenant=%s name=%s", principal.tenant_id, body.name.strip())
    return {"ok": True, "tenant_name": body.name.strip()}


# ── Routes: documents ─────────────────────────────────────────────────────────
@app.get("/api/documents", response_model=list[DocumentMeta])
def list_documents(principal: Annotated[Principal, Depends(principal_from_auth)]) -> list[DocumentMeta]:
    docs = index_store.list_docs(principal.tenant_id)
    return [
        DocumentMeta(
            file_name=d["file_name"],
            chunks=len(d.get("chunks", d.get("pages", []))),
            uploaded_at=d["uploaded_at"],
        )
        for d in docs
    ]


@app.delete("/api/documents/{file_name}")
def delete_document(
    file_name: str,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    require_role(principal, "admin")
    tenant_id = principal.tenant_id
    safe_name = safe_filename(file_name)
    deleted = index_store.delete_doc(tenant_id, safe_name)
    if not deleted:
        raise HTTPException(status_code=404, detail="Document not found")
    storage.delete(tenant_id, safe_name)
    log.info("deleted document tenant=%s file=%s", tenant_id, safe_name)
    return {"ok": True, "file_name": safe_name}


# ── Routes: upload ────────────────────────────────────────────────────────────
@app.post("/api/upload")
async def upload_document(
    request: Request,
    principal: Annotated[Principal, Depends(principal_from_auth)],
    file: UploadFile = File(...),
) -> dict:
    require_role(principal, "admin", "uploader")
    tenant_id = principal.tenant_id

    content_length = request.headers.get("content-length")
    if content_length and int(content_length) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="Upload exceeds 50 MB limit")

    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="Upload exceeds 50 MB limit")
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename required")

    file_name = safe_filename(file.filename)
    chunks = extract_chunks(content, file_name)  # PDF or .docx; rejects others
    uri = storage.upload(tenant_id, file_name, content)

    # Embed each chunk so it is retrievable by vector similarity
    try:
        vectors = embedder.embed([c["text"] for c in chunks])
        for c, v in zip(chunks, vectors):
            c["embedding"] = v
    except Exception as exc:
        log.warning("embedding failed for %s (%s); stored without vectors", file_name, exc)

    index_store.upsert_doc(tenant_id, {
        "tenant_id": tenant_id,
        "file_name": file_name,
        "storage_uri": uri,
        "chunks": chunks,
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    })
    log.info("upload tenant=%s file=%s chunks=%d uri=%s", tenant_id, file_name, len(chunks), uri)
    return {"tenant_id": tenant_id, "file_name": file_name, "chunks": len(chunks)}


# ── Routes: chat ──────────────────────────────────────────────────────────────
@app.post("/api/chat", response_model=ChatResponse)
def chat(req: ChatRequest, principal: Annotated[Principal, Depends(principal_from_auth)]) -> ChatResponse:
    tenant_id = principal.tenant_id
    queries_used = enforce_quota(tenant_id)
    history = [t.model_dump() for t in req.history]

    started = time.perf_counter()
    docs = index_store.list_docs(tenant_id)
    total_chunks = sum(len(d.get("chunks") or d.get("pages", [])) for d in docs)

    # 1) Query rewriting — expand/resolve the question for better retrieval
    search_query = req.message
    if hasattr(generator, "rewrite_query"):
        try:
            search_query = generator.rewrite_query(req.message, history)
            if search_query != req.message:
                log.info("query_rewrite tenant=%s original=%r rewritten=%r", tenant_id, req.message, search_query)
        except Exception as exc:
            log.warning("query rewrite failed (%s); using original", exc)

    query_terms = tokenize(search_query)
    denom = max(len(query_terms), 1)

    # 2) Vector retrieval (top candidates by cosine)
    candidates: list[dict] = []
    try:
        q_vec = embedder.embed_query(search_query)
        candidates = index_store.vector_search(tenant_id, q_vec, k=30)
    except Exception as exc:
        log.warning("vector search failed (%s); keyword fallback", exc)

    # 3) Fallback: if no embedded chunks, scan all chunks (keyword only)
    if not candidates:
        for doc in docs:
            chunks = doc.get("chunks") or [
                {"page": i + 1, "chunk_index": 0, "text": p}
                for i, p in enumerate(doc.get("pages", []))
            ]
            for ch in chunks:
                candidates.append({
                    "file_name": doc["file_name"],
                    "page": ch.get("page", 1),
                    "chunk_index": ch.get("chunk_index", 0),
                    "text": ch.get("text", ""),
                    "vec_score": 0.0,
                })

    # 4) Hybrid score: blend cosine with keyword overlap
    ranked: list[tuple[float, dict]] = []
    for c in candidates:
        kw = len(query_terms & tokenize(c["text"])) / denom
        vec = float(c.get("vec_score", 0.0))
        score = HYBRID_VECTOR_WEIGHT * vec + (1 - HYBRID_VECTOR_WEIGHT) * kw
        if score > 0:
            ranked.append((round(score, 4), c))

    ranked.sort(key=lambda t: t[0], reverse=True)
    # Drop chunks below relevance floor to avoid poisoning Gemini with noise
    ranked = [(s, c) for s, c in ranked if s >= 0.05]
    top = ranked[:6]
    top_chunks = [c for _, c in top]

    answer = generator.generate(req.message, top_chunks, history)

    citations = [
        Citation(
            file_name=c["file_name"],
            page=c.get("page", 1),
            chunk_index=c.get("chunk_index", 0),
            excerpt=c["text"][:280].replace("\n", " "),
            score=s,
        )
        for s, c in top
    ]
    latency_ms = max(int((time.perf_counter() - started) * 1000), 1)

    retrieval = RetrievalTrace(
        engine=RAG_ENGINE_LABEL,
        query_terms=sorted(query_terms),
        chunks_searched=total_chunks,
        candidates_ranked=len(ranked),
        top_k=len(citations),
        max_score=top[0][0] if top else 0.0,
        latency_ms=latency_ms,
    )

    # Record for knowledge analytics + gap detection (answer stored when confidence is high)
    insights_store.record(tenant_id, req.message, retrieval.max_score, answer=answer)

    log.info(
        "chat tenant=%s terms=%d searched=%d ranked=%d matches=%d latency_ms=%d queries_used=%d",
        tenant_id, len(query_terms), total_chunks, len(ranked), len(citations), latency_ms, queries_used,
    )
    return ChatResponse(
        answer=answer,
        citations=citations,
        retrieval=retrieval,
        tenant_id=tenant_id,
        queries_used=queries_used,
        query_limit=QUERY_LIMIT,
    )


# ── Routes: insights (knowledge analytics + gap detection) ────────────────────
@app.get("/api/insights", response_model=InsightsResponse)
def get_insights(principal: Annotated[Principal, Depends(principal_from_auth)]) -> InsightsResponse:
    require_role(principal, "admin")
    data = insights_store.summary(principal.tenant_id)
    return InsightsResponse(**data)


# ── Routes: members ───────────────────────────────────────────────────────────
@app.get("/api/tenant/members", response_model=list[MemberOut])
def list_members(principal: Annotated[Principal, Depends(principal_from_auth)]) -> list[MemberOut]:
    require_role(principal, "admin")
    members = member_store.get_members(principal.tenant_id)
    return [MemberOut(uid=m.uid, email=m.email, role=m.role, invited_at=m.invited_at, joined_at=m.joined_at)
            for m in members]


@app.post("/api/tenant/invite", status_code=201)
def invite_member(
    body: InviteRequest,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    require_role(principal, "admin")
    if body.role not in VALID_ROLES:
        raise HTTPException(status_code=400, detail=f"Invalid role. Must be one of {list(VALID_ROLES)}")
    tenant_id = principal.tenant_id
    now = datetime.now(timezone.utc).isoformat()
    uid = f"invited-{re.sub(r'[^a-z0-9]', '-', body.email.lower())}"

    # Prevent duplicate email
    existing = member_store.get_members(tenant_id)
    if any(m.email == body.email for m in existing):
        raise HTTPException(status_code=409, detail="Member with this email already exists")

    member = MemberRecord(uid=uid, email=body.email, role=body.role, invited_at=now, joined_at=None)
    member_store.add_member(tenant_id, member)

    invite_link = f"{config.app_base_url}/?invite={uid}&tenant={tenant_id}"
    tenant_name = tenant_profile_store.get_name(tenant_id) or prettify(tenant_id)
    inviter = principal.email or "A teammate"

    # Send a real email if SMTP is configured; otherwise return the link so
    # the admin can share it manually (dev / unconfigured prod).
    email_sent = send_invite_email(
        to_email=body.email, invite_link=invite_link,
        tenant_name=tenant_name, inviter=inviter, role=body.role,
    )
    log.info("invite tenant=%s email=%s role=%s emailed=%s", tenant_id, body.email, body.role, email_sent)
    return {"ok": True, "uid": uid, "invite_link": invite_link, "email_sent": email_sent}


@app.post("/api/tenant/accept-invite")
def accept_invite(
    body: AcceptInviteRequest,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    """Join a workspace the caller was invited to. Authorization is by the
    caller's *verified email* matching a pending invite in the target tenant —
    so a leaked link can't be claimed by someone else's account."""
    target = body.tenant_id
    email = (principal.email or "").strip().lower()
    if not email:
        raise HTTPException(status_code=403, detail="Your account has no email; cannot accept an invite")

    members = member_store.get_members(target)

    # Already a member (re-click / refresh) → idempotent success.
    already = next((m for m in members if m.uid == principal.uid), None)
    if already and already.joined_at:
        set_tenant_claim(principal.uid, target)  # ensure claim is present
        return {"ok": True, "tenant_id": target, "role": already.role, "already_member": True}

    pending = next(
        (m for m in members if (m.email or "").strip().lower() == email and not m.joined_at),
        None,
    )
    if pending is None:
        raise HTTPException(status_code=404, detail="No pending invite for your email in this workspace")

    now = datetime.now(timezone.utc).isoformat()
    role = pending.role
    # Swap the placeholder invite row for the caller's real uid.
    if pending.uid != principal.uid:
        member_store.remove_member(target, pending.uid)
    member_store.add_member(target, MemberRecord(
        uid=principal.uid, email=principal.email or email, role=role,
        invited_at=pending.invited_at, joined_at=now,
    ))

    # Persist the shared tenant on the user so future tokens resolve to it.
    claim_set = set_tenant_claim(principal.uid, target)
    log.info("accept-invite tenant=%s uid=%s role=%s claim=%s", target, principal.uid, role, claim_set)
    return {"ok": True, "tenant_id": target, "role": role, "claim_set": claim_set}


@app.patch("/api/tenant/members/{uid}")
def update_member_role(
    uid: str,
    body: RoleRequest,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    require_role(principal, "admin")
    if body.role not in VALID_ROLES:
        raise HTTPException(status_code=400, detail=f"Invalid role. Must be one of {list(VALID_ROLES)}")
    tenant_id = principal.tenant_id
    members = member_store.get_members(tenant_id)
    if not any(m.uid == uid for m in members):
        raise HTTPException(status_code=404, detail="Member not found")
    member_store.update_role(tenant_id, uid, body.role)
    log.info("role_change tenant=%s uid=%s role=%s by=%s", tenant_id, uid, body.role, principal.uid)
    return {"ok": True, "uid": uid, "role": body.role}


@app.delete("/api/tenant/members/{uid}")
def remove_member(
    uid: str,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    require_role(principal, "admin")
    tenant_id = principal.tenant_id
    # Prevent self-removal
    if uid == principal.uid:
        raise HTTPException(status_code=400, detail="Cannot remove yourself")
    members = member_store.get_members(tenant_id)
    if not any(m.uid == uid for m in members):
        raise HTTPException(status_code=404, detail="Member not found")
    member_store.remove_member(tenant_id, uid)
    log.info("remove_member tenant=%s uid=%s by=%s", tenant_id, uid, principal.uid)
    return {"ok": True, "uid": uid}


# ── Routes: tenant hard-erase ─────────────────────────────────────────────────
@app.delete("/api/tenant")
def delete_tenant(principal: Annotated[Principal, Depends(principal_from_auth)]) -> dict:
    require_role(principal, "admin")
    tenant_id = principal.tenant_id

    # Delete all files and index entries for this tenant
    storage.delete_tenant(tenant_id)
    removed = index_store.delete_tenant(tenant_id)

    # Clear usage counters, members, and profile
    usage_store.reset_tenant(tenant_id)
    member_store.delete_tenant(tenant_id)
    tenant_profile_store.delete_tenant(tenant_id)

    log.info("tenant_erased tenant=%s docs=%d by=%s", tenant_id, removed, principal.uid)
    return {"ok": True, "tenant_id": tenant_id, "documents_removed": removed}


# ── Routes: dev reset ─────────────────────────────────────────────────────────
@app.post("/api/dev/reset")
def reset_dev_state() -> dict:
    if config.env == "production":
        raise HTTPException(status_code=403, detail="Not available in production")
    usage_store.reset()
    insights_store.reset()
    log.info("quota + insights reset by dev endpoint")
    return {"ok": True}


# ── Routes: share ─────────────────────────────────────────────────────────────
@app.post("/api/share", status_code=201)
def create_share(
    body: ShareCreateRequest,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    share_id = share_store.create({
        "question": body.question,
        "answer": body.answer,
        "citations": body.citations,
        "tenant_id": principal.tenant_id,
    })
    url = f"{config.app_base_url}/share/{share_id}"
    log.info("share_create tenant=%s share_id=%s", principal.tenant_id, share_id)
    return {"share_id": share_id, "url": url}


@app.get("/api/share/{share_id}")
def get_share(share_id: str) -> dict:
    data = share_store.get(share_id)
    if not data:
        raise HTTPException(status_code=404, detail="Share not found")
    return data


# ── Routes: widget keys ───────────────────────────────────────────────────────
@app.post("/api/tenant/widget-keys", status_code=201)
def create_widget_key(
    body: WidgetKeyCreateRequest,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    require_role(principal, "admin")
    result = widget_key_store.create_key(principal.tenant_id, body.label)
    log.info("widget_key_created tenant=%s key_id=%s label=%s", principal.tenant_id, result["key_id"], body.label)
    return result  # raw key returned ONCE — not stored in plaintext after this


@app.get("/api/tenant/widget-keys")
def list_widget_keys(principal: Annotated[Principal, Depends(principal_from_auth)]) -> list[dict]:
    require_role(principal, "admin")
    return widget_key_store.list_keys(principal.tenant_id)


@app.delete("/api/tenant/widget-keys/{key_id}")
def delete_widget_key(
    key_id: str,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    require_role(principal, "admin")
    deleted = widget_key_store.delete_key(principal.tenant_id, key_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Widget key not found")
    log.info("widget_key_deleted tenant=%s key_id=%s", principal.tenant_id, key_id)
    return {"ok": True, "key_id": key_id}


# ── Routes: Slack integration ─────────────────────────────────────────────────
@app.post("/api/integrations/slack/connect")
def slack_connect(
    body: SlackConnectRequest,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    require_role(principal, "admin")
    slack_store.connect(body.team_id, body.team_name, principal.tenant_id)
    slash_cmd = "/ask"
    log.info("slack_connect tenant=%s team=%s", principal.tenant_id, body.team_id)
    return {
        "ok": True,
        "team_id": body.team_id,
        "instructions": (
            f"Add a Slack slash command:\n"
            f"  Command: {slash_cmd}\n"
            f"  Request URL: {config.app_base_url.replace('web.app', 'run.app') if 'web.app' in config.app_base_url else config.app_base_url}/api/integrations/slack/command\n"
            f"  Short description: Ask your company knowledge base\n"
            f"Set SLACK_SIGNING_SECRET env var on Cloud Run to the signing secret from your Slack app."
        ),
    }


@app.get("/api/integrations/slack/connections")
def list_slack_connections(principal: Annotated[Principal, Depends(principal_from_auth)]) -> list[dict]:
    require_role(principal, "admin")
    return slack_store.list_connections(principal.tenant_id)


@app.delete("/api/integrations/slack/connections/{team_id}")
def slack_disconnect(
    team_id: str,
    principal: Annotated[Principal, Depends(principal_from_auth)],
) -> dict:
    require_role(principal, "admin")
    deleted = slack_store.disconnect(team_id, principal.tenant_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Slack connection not found")
    return {"ok": True, "team_id": team_id}


def _verify_slack_signature(body: bytes, timestamp: str, signature: str) -> bool:
    if not SLACK_SIGNING_SECRET:
        return False
    # Replay attack prevention: reject requests older than 5 minutes
    try:
        if abs(time.time() - float(timestamp)) > 300:
            return False
    except ValueError:
        return False
    base = f"v0:{timestamp}:{body.decode()}"
    expected = "v0=" + hmac.new(
        SLACK_SIGNING_SECRET.encode(), base.encode(), hashlib.sha256
    ).hexdigest()
    return hmac.compare_digest(expected, signature)


def _run_slack_rag(tenant_id: str, question: str, user_name: str, response_url: str) -> None:
    """Background task: run RAG and POST result to Slack response_url."""
    try:
        docs = index_store.list_docs(tenant_id)
        total_chunks = sum(len(d.get("chunks") or d.get("pages", [])) for d in docs)
        if not docs:
            answer = "No documents have been uploaded to your knowledge base yet."
        else:
            search_query = question
            history: list[dict] = []
            if hasattr(generator, "rewrite_query"):
                try:
                    search_query = generator.rewrite_query(question, history)
                except Exception:
                    pass

            query_terms = tokenize(search_query)
            denom = max(len(query_terms), 1)
            candidates: list[dict] = []
            try:
                q_vec = embedder.embed_query(search_query)
                candidates = index_store.vector_search(tenant_id, q_vec, k=30)
            except Exception:
                pass

            if not candidates:
                for doc in docs:
                    chunks = doc.get("chunks") or [
                        {"page": i + 1, "chunk_index": 0, "text": p}
                        for i, p in enumerate(doc.get("pages", []))
                    ]
                    for ch in chunks:
                        candidates.append({
                            "file_name": doc["file_name"], "page": ch.get("page", 1),
                            "chunk_index": ch.get("chunk_index", 0),
                            "text": ch.get("text", ""), "vec_score": 0.0,
                        })

            ranked = []
            for c in candidates:
                kw = len(query_terms & tokenize(c["text"])) / denom
                vec = float(c.get("vec_score", 0.0))
                score = HYBRID_VECTOR_WEIGHT * vec + (1 - HYBRID_VECTOR_WEIGHT) * kw
                if score >= 0.05:
                    ranked.append((round(score, 4), c))
            ranked.sort(key=lambda t: t[0], reverse=True)
            top_chunks = [c for _, c in ranked[:6]]
            answer = generator.generate(question, top_chunks, [])
            max_score = ranked[0][0] if ranked else 0.0
            insights_store.record(tenant_id, question, max_score, answer=answer)

        # Build citation footnotes
        sources = ""
        if top_chunks if 'top_chunks' in dir() else False:
            seen = []
            for c in (top_chunks if 'top_chunks' in dir() else []):
                ref = f"{c['file_name']} p.{c.get('page', 1)}"
                if ref not in seen:
                    seen.append(ref)
            if seen:
                sources = "\n_Sources: " + " · ".join(seen) + "_"

        payload = json.dumps({
            "response_type": "in_channel",
            "text": f"*{user_name} asked:* {question}\n\n{answer}{sources}",
        }).encode()
        req = urllib.request.Request(
            response_url, data=payload,
            headers={"Content-Type": "application/json"}, method="POST",
        )
        urllib.request.urlopen(req, timeout=10)
    except Exception as exc:
        log.warning("slack_rag failed tenant=%s err=%s", tenant_id, exc)
        try:
            payload = json.dumps({"response_type": "ephemeral", "text": "⚠️ Something went wrong. Try again."}).encode()
            req = urllib.request.Request(response_url, data=payload,
                                          headers={"Content-Type": "application/json"}, method="POST")
            urllib.request.urlopen(req, timeout=5)
        except Exception:
            pass


@app.post("/api/integrations/slack/command")
async def slack_command(request: Request, background_tasks: BackgroundTasks) -> dict:
    """Slack slash command handler — no Firebase auth, uses Slack signing secret."""
    body = await request.body()

    ts = request.headers.get("X-Slack-Request-Timestamp", "")
    sig = request.headers.get("X-Slack-Signature", "")
    if SLACK_SIGNING_SECRET and not _verify_slack_signature(body, ts, sig):
        raise HTTPException(status_code=401, detail="Invalid Slack signature")

    form = await request.form()
    team_id   = str(form.get("team_id", ""))
    text      = str(form.get("text", "")).strip()
    user_name = str(form.get("user_name", "someone"))
    response_url = str(form.get("response_url", ""))

    tenant_id = slack_store.get_tenant(team_id)
    if not tenant_id:
        return {
            "response_type": "ephemeral",
            "text": "❌ This Slack workspace isn't connected to RAGaaS. An admin must connect it first at your RAGaaS dashboard.",
        }
    if not text:
        return {"response_type": "ephemeral", "text": "Usage: `/ask [your question]`"}

    log.info("slack_command tenant=%s team=%s user=%s question=%r", tenant_id, team_id, user_name, text)
    background_tasks.add_task(_run_slack_rag, tenant_id, text, user_name, response_url)
    return {
        "response_type": "in_channel",
        "text": f"⏳ *{user_name}* asked: _{text}_\nSearching your knowledge base…",
    }
